"""Document generators - Logos scaled to 50% of original size"""

from io import BytesIO
from typing import List, Dict
import os

try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except:
    WEASYPRINT_AVAILABLE = False

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils.company_template import PARTNER_LOGO_URLS, PARTNER_LOGOS

def get_cover_page_html(title: str, topic: str, partner_institute: str) -> str:
    logo_url = PARTNER_LOGO_URLS.get(partner_institute, PARTNER_LOGO_URLS["Default"])
    return f"""
    <div class="cover-page">
        <div class="cover-main">
            <div class="cover-content">
                <h1 class="cover-title">{title}</h1>
                <h2 class="cover-topic">{topic}</h2>
            </div>
        </div>
        <div class="cover-footer">
            <div class="footer-logo-wrapper">
                <img src="{logo_url}" class="partner-banner" alt="{partner_institute}">
            </div>
        </div>
    </div>
    """

class PDFGenerator:
    def __init__(self, title: str, topic: str, partner_institute: str = "IIT Kanpur"):
        self.title = title
        self.topic = topic
        self.partner_institute = partner_institute
    
    def generate(self, qa_pairs: List[Dict]) -> bytes:
        if not WEASYPRINT_AVAILABLE:
            raise Exception("WeasyPrint not available")
        
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page cover {{
            size: A4;
            margin: 0;
        }}
        
        @page content {{
            size: A4;
            margin: 18.3mm 18.3mm 18.3mm 18.3mm;
        }}

        body {{
            margin: 0;
            padding: 0;
            font-family: Calibri, sans-serif;
            font-size: 14pt;
            line-height: 1.5;
            color: #000000;
        }}
        
        /* COVER PAGE */
        .cover-page {{
            page: cover;
            height: 297mm;
            width: 210mm;
            display: flex;
            flex-direction: column;
            justify-content: flex-end;
            page-break-after: always;
        }}
        
        .cover-main {{
            background-color: #3030ff;
            flex: 1;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            min-height: 87%;
            height: 87%;
            padding: 0;
        }}
        .cover-content {{
            width: 100%;
            height: 100%;
            text-align: center;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            gap: 0;
        }}
        .cover-title {{
            font-size: 27pt;
            font-weight: bold;
            color: #FFFFFF;
            font-family: Calibri, sans-serif;
            margin: 0;
            padding: 0;
            line-height: 1.2;
        }}
        
        .cover-topic {{
            font-size: 27pt;
            font-weight: normal;
            color: #FFFFFF;
            font-family: Calibri, sans-serif;
            margin: 0;
            padding: 0;
            line-height: 1.2;
        }}
        .cover-footer {{
            height: 13%;
            background: #FFF;
            width: 100%;
            min-height: 100px;
            display: flex;
            justify-content: center;
            align-items: flex-end;
            position: relative;
            text-align: center;
            padding: 10px 0;
            margin: 0;
        }}
        .footer-logo-wrapper {{
            width: 100%;
            height: 100%;
            display: flex;
            justify-content: center;
            align-items: center;
        }}
        .partner-banner {{
            /* Scale the logo to 50% of its natural size */
            zoom: 0.5; 
            max-width: 100%;
            height: auto;
            width: auto;
            display: block;
            margin: 0 auto;
            padding: 0;
        }}

        /* CONTENT PAGES */
        .content-page {{
            page: content;
        }}

        .question-block {{
            margin-bottom: 15px;
        }}

        .question-header {{
            font-size: 14pt;
            font-weight: bold;
            margin-bottom: 5px;
            color: #000000;
            font-family: Calibri, sans-serif;
            text-align: justify;
        }}

        .answer-header {{
            font-size: 14pt;
            font-weight: bold;
            color: #000000;
            font-family: Calibri, sans-serif;
            display: inline;
            text-align: justify;
        }}

        .answer-text {{
            font-size: 14pt;
            text-align: justify;
            line-height: 1.5;
            margin-bottom: 10px;
            color: #000000;
            font-family: Calibri, sans-serif;
            display: inline;
        }}
    </style>
</head>
<body>
{get_cover_page_html(self.title, self.topic, self.partner_institute)}
<div class="content-page">
"""
        for i, qa in enumerate(qa_pairs, 1):
            question = qa.get('question', '')
            answer = qa.get('answer', '')
            html_content += f"""
<div class="question-block">
    <div class="question-header">Question {i}: {question}</div>
    <div class="answer-header">Answer: </div><div class="answer-text">{answer}</div>
</div>
"""
        html_content += """
</div>
</body>
</html>"""
        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes

class WordDocumentGenerator:
    def __init__(self):
        pass
    
    def _set_cell_background(self, cell, color):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def generate(self, qa_pairs: List[Dict], title: str, topic: str, partner_institute: str = "IIT Kanpur") -> bytes:
        doc = Document()
        
        # === FIRST PAGE - COVER PAGE WITH BLUE BACKGROUND ===
        # Set margins to 0 for full-page blue background
        section = doc.sections[0]
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        
        # Create a table for full-page background (1 row, 1 column)
        cover_table = doc.add_table(rows=1, cols=1)
        cover_table.autofit = False
        cover_table.allow_autofit = False
        
        # Set table width to full page
        cover_table.width = Inches(8.5)
        cover_cell = cover_table.rows[0].cells[0]
        cover_cell.width = Inches(8.5)
        
        # Set cell background to blue
        self._set_cell_background(cover_cell, '3030FF')
        
        # Set cell margins to 0
        tcPr = cover_cell._element.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), '0')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
        
        # Add title and topic to the blue cell
        title_para = cover_cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(100)
        title_run = title_para.add_run(title)
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(27)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Add topic
        topic_para = cover_cell.add_paragraph()
        topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        topic_para.paragraph_format.space_after = Pt(100)
        topic_run = topic_para.add_run(topic)
        topic_run.font.name = 'Calibri'
        topic_run.font.size = Pt(27)
        topic_run.font.bold = False
        topic_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Add logo at bottom of cell
        logo_para = cover_cell.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.paragraph_format.space_before = Pt(50)
        
        logo_path = PARTNER_LOGOS.get(partner_institute, PARTNER_LOGOS["Default"])
        if os.path.exists(logo_path):
            try:
                run = logo_para.add_run()
                picture = run.add_picture(logo_path, width=Inches(2.0))
                # Scale to 50%
                picture.width = int(picture.width * 0.5)
                picture.height = int(picture.height * 0.5)
            except Exception as e:
                pass
        
        # Remove table borders
        tbl = cover_table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # === PAGE BREAK ===
        doc.add_page_break()
        
        # === CONTENT PAGES ===
        section = doc.sections[-1]
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
        
        # Add Q&A content
        for i, qa in enumerate(qa_pairs, 1):
            question = qa.get('question', '')
            answer = qa.get('answer', '')
            
            # Question - 14pt bold, justified
            q_para = doc.add_paragraph()
            q_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            q_run = q_para.add_run(f"Question {i}: {question}")
            q_run.font.name = 'Calibri'
            q_run.font.size = Pt(14)
            q_run.font.bold = True
            q_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            q_para.paragraph_format.space_after = Pt(6)
            
            # Answer - 14pt, inline, justified
            ans_para = doc.add_paragraph()
            ans_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # "Answer:" - 14pt bold
            ans_header_run = ans_para.add_run("Answer: ")
            ans_header_run.font.name = 'Calibri'
            ans_header_run.font.size = Pt(14)
            ans_header_run.font.bold = True
            
            # Answer text - 14pt normal
            ans_run = ans_para.add_run(answer)
            ans_run.font.name = 'Calibri'
            ans_run.font.size = Pt(14)
            ans_run.font.bold = False
            
            ans_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            ans_para.paragraph_format.space_after = Pt(10)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
